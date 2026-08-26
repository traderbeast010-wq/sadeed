#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
استخراج نصّ العقد من الملف المرفوع — بلا نموذج.

PDFs العقود ليست PDFs القوانين: الأخيرة تستخدم خطوطاً رسمية قديمة بترميز
تالف، أمّا العقود فتُكتب غالباً في Word وتُصدَّر، فترميزها سليم. لذلك نستخدم
PyMuPDF هنا (أدقّ في الحفاظ على ترتيب الفقرات) مع تمرير النصّ على نفس دالة
الإصلاح احتياطاً — فهي غير مؤذية للنصّ السليم.
"""

import os
import re
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ROOT, "corpus"))
from repair import repair  # noqa: E402


class ExtractionError(Exception):
    """يُرفع حين لا يمكن استخراج نصّ قابل للقراءة."""


MIN_CHARS = 120          # أقلّ من هذا: الملف صورة أو فارغ


def from_pdf(path: str) -> str:
    try:
        import pymupdf
    except ImportError:
        import fitz as pymupdf
    doc = pymupdf.open(path)
    try:
        pages = [p.get_text("text") for p in doc]
    finally:
        doc.close()
    return "\n".join(pages)


def from_docx(path: str) -> str:
    from docx import Document
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    # الجداول شائعة في العقود (بيانات الأطراف، جداول الأجور)
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract(path: str) -> str:
    if not os.path.exists(path):
        raise ExtractionError(f"الملف غير موجود: {path}")

    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        raw = from_pdf(path)
    elif ext in (".docx", ".doc"):
        raw = from_docx(path)
    elif ext == ".txt":
        with open(path, encoding="utf-8") as f:
            raw = f.read()
    else:
        raise ExtractionError(f"صيغة غير مدعومة «{ext}» — يُقبل PDF أو Word فقط")

    text = repair(raw)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if len(text) < MIN_CHARS:
        raise ExtractionError(
            "الملف لا يحتوي نصاً قابلاً للقراءة. غالباً صورة ممسوحة ضوئياً — "
            "النسخة الحالية لا تدعم التعرّف الضوئي على الحروف."
        )
    arabic = sum(1 for c in text if "ء" <= c <= "ي")
    if arabic / max(len(text), 1) < 0.15:
        raise ExtractionError("الملف لا يبدو عقداً عربياً.")
    return text
